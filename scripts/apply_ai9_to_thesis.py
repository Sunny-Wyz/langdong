#!/usr/bin/env python3
"""Inject ai9 experiment content into thesis chapter 5.3 and output a new docx."""

from __future__ import annotations

import math
from pathlib import Path

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


INPUT_DOC = Path("/Users/weiyaozhou/Desktop/test/备件管理系统 完整论文（修订版）_副本_已修改.docx")
OUTPUT_DOC = Path("/Users/weiyaozhou/Desktop/test/备件管理系统 完整论文（修订版）_副本_已修改_ai9版.docx")
IMG_RF = Path("/Users/weiyaozhou/Desktop/test/rf_vs_avg_5_25.png")
IMG_SS = Path("/Users/weiyaozhou/Desktop/test/ss_vs_fixed_5_26.png")


AI9_INPUT_ROWS = [
    ["2025-01", 30, 19, 14, 24],
    ["2025-02", 10, 21, 16, 26],
    ["2025-03", 31, 19, 14, 24],
    ["2025-04", 11, 21, 16, 26],
    ["2025-05", 29, 19, 14, 24],
    ["2025-06", 9, 21, 16, 26],
    ["2025-07", 32, 19, 14, 24],
    ["2025-08", 12, 21, 16, 26],
    ["2025-09", 30, 19, 14, 24],
    ["2025-10", 10, 21, 16, 26],
    ["2025-11", 31, 19, 14, 24],
    ["2025-12", 11, 21, 16, 26],
]

RF_ROWS = [
    ["2025-07", 32.00, 29.00, 20.00, 3.00, 12.00, "9.38%", "37.50%", "是"],
    ["2025-08", 12.00, 9.00, 21.71, 3.00, 9.71, "25.00%", "80.95%", "是"],
    ["2025-09", 30.00, 32.00, 20.50, 2.00, 9.50, "6.67%", "31.67%", "是"],
    ["2025-10", 10.00, 12.00, 21.56, 2.00, 11.56, "20.00%", "115.56%", "是"],
    ["2025-11", 31.00, 30.00, 20.40, 1.00, 10.60, "3.23%", "34.19%", "是"],
    ["2025-12", 11.00, 10.00, 21.36, 1.00, 10.36, "9.09%", "94.21%", "是"],
]

SS_SUMMARY_ROWS = [
    ["动态法", "100.00%", "0.00", "13.00"],
    ["固定法（SS=10）", "33.33%", "1.08", "10.00"],
]


def setup_font() -> None:
    # Try common CJK fonts; matplotlib will pick the first available.
    plt.rcParams["font.sans-serif"] = [
        "PingFang SC",
        "Heiti SC",
        "STHeiti",
        "Arial Unicode MS",
        "SimHei",
        "DejaVu Sans",
    ]
    plt.rcParams["axes.unicode_minus"] = False


def generate_images() -> None:
    setup_font()

    months = ["2025-07", "2025-08", "2025-09", "2025-10", "2025-11", "2025-12"]
    rf_err = [3, 3, 2, 2, 1, 1]
    avg_err = [12, 9.71, 9.50, 11.56, 10.60, 10.36]

    x = range(len(months))
    width = 0.38

    plt.figure(figsize=(10, 4.8))
    plt.bar([i - width / 2 for i in x], rf_err, width=width, label="RF绝对误差")
    plt.bar([i + width / 2 for i in x], avg_err, width=width, label="简单平均法绝对误差")
    plt.xticks(list(x), months)
    plt.ylabel("绝对误差")
    plt.title("图5-25  RF与简单平均法绝对误差对比")
    plt.legend()
    plt.tight_layout()
    plt.savefig(IMG_RF, dpi=200)
    plt.close()

    dims = ["覆盖率(%)", "平均缺口", "平均安全库存"]
    dyn = [100.00, 0.00, 13.00]
    fix = [33.33, 1.08, 10.00]
    x2 = range(len(dims))

    plt.figure(figsize=(8.6, 4.8))
    plt.bar([i - width / 2 for i in x2], dyn, width=width, label="动态法")
    plt.bar([i + width / 2 for i in x2], fix, width=width, label="固定法")
    plt.xticks(list(x2), dims)
    plt.ylabel("数值")
    plt.title("图5-26  动态安全库存与固定法对比")
    plt.legend()
    plt.tight_layout()
    plt.savefig(IMG_SS, dpi=200)
    plt.close()


def add_table_before(anchor, doc: Document, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        # Template may not include this English style name.
        pass
    for i, h in enumerate(headers):
        table.cell(0, i).text = str(h)
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            table.cell(r_idx, c_idx).text = str(val)
    anchor._p.addprevious(table._tbl)
    return table


def replace_placeholder_with_image(doc: Document, marker: str, image_path: Path, width_in: float = 5.8) -> bool:
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = "\n".join(p.text for p in cell.paragraphs).strip()
                if marker in text:
                    # clear existing paragraphs
                    for p in list(cell.paragraphs):
                        p._element.getparent().remove(p._element)
                    p = cell.add_paragraph()
                    run = p.add_run()
                    run.add_picture(str(image_path), width=Inches(width_in))
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    return True
    return False


def insert_ai9_section(doc: Document) -> None:
    paragraphs = doc.paragraphs
    algo2_idx = next((i for i, p in enumerate(paragraphs) if "Algorithm 2" in p.text), None)
    if algo2_idx is None:
        raise RuntimeError("未找到 Algorithm 2 锚点。")

    summary_idx = next((i for i in range(algo2_idx + 1, len(paragraphs)) if paragraphs[i].text.strip() == "本章小结"), None)
    if summary_idx is None:
        raise RuntimeError("未找到 Algorithm 2 后的 本章小结 锚点。")

    anchor = paragraphs[summary_idx]

    def add_para(text: str):
        p = anchor.insert_paragraph_before(text)
        return p

    add_para("5.3.3 基于ai9数据的补充验证（SP20001）")
    add_para("（1）数据来源与样本说明")
    add_para(
        "本补充实验采用ai9可控造数数据，以SP20001深沟球轴承为样本，覆盖2025年1月至12月。"
        "其中2025-07至2025-12作为滚动评估窗口，用于比较随机森林（RF）与简单平均法。"
    )
    add_para("表5-5  ai9输入数据（2025-01~2025-12）")
    add_table_before(
        anchor,
        doc,
        ["月份", "实际需求", "预测需求", "预测下界", "预测上界"],
        AI9_INPUT_ROWS,
    )
    add_para("参数意义：实际需求为真实值；预测需求及上下界用于构建RF特征与动态安全库存计算。")

    add_para("（2）RF与简单平均法对比结果")
    add_para("表5-6  RF与简单平均法逐月对比（2025-07~2025-12）")
    add_table_before(
        anchor,
        doc,
        ["月份", "实际需求", "RF预测", "简单平均预测", "RF绝对误差", "平均法绝对误差", "RF相对误差", "平均法相对误差", "RF是否更优"],
        RF_ROWS,
    )
    add_para("参数意义：RF是否更优=是，表示该月RF绝对误差小于简单平均法。")

    add_para("简单平均法预测定义为：y_hat_t(avg) = (1/(t-1)) * Σ(i=1..t-1) y_i。                    （5.5）")
    add_para("MAE = (1/n) * Σ|y_t - y_hat_t|。                                                 （5.6）")
    add_para("RMSE = sqrt((1/n) * Σ(y_t - y_hat_t)^2)。                                       （5.7）")
    add_para("MAPE = (1/n) * Σ(|y_t - y_hat_t|/y_t) * 100%。                                  （5.8）")
    add_para(
        "以2025-07为例：简单平均预测=(30+10+31+11+29+9)/6=20.00，RF预测=29.00，"
        "对应绝对误差分别为12.00与3.00。"
    )

    add_para("（3）动态安全库存与固定法对比结果")
    add_para("动态安全库存计算采用：σ_d=(upper_bound-lower_bound)/(2*1.645)。                    （5.9）")
    add_para("SS_dynamic=ceil(k*σ_d*sqrt(L))，其中k=1.28，L=10。                                 （5.10）")
    add_para("ROP_dynamic=ceil((forecast_qty/30)*L + SS_dynamic)。                               （5.11）")
    add_para("覆盖判定：covered=1{SS>=|actual-forecast|}，缺口：shortage=max(0,|actual-forecast|-SS)。 （5.12）")
    add_para("表5-7  安全库存方法汇总指标对比")
    add_table_before(
        anchor,
        doc,
        ["方法", "覆盖率", "平均缺口", "平均安全库存"],
        SS_SUMMARY_ROWS,
    )
    add_para("参数意义：覆盖率与平均缺口是优劣判断核心指标；平均安全库存用于反映库存占用。")

    add_para("（4）补充实验结论")
    add_para(
        "在ai9补充实验中，RF在评估窗口6/6个月均优于简单平均法（MAE=2.00，RMSE=2.16，MAPE=12.23%）；"
        "动态安全库存法覆盖率100.00%、平均缺口0.00，优于固定法（覆盖率33.33%、平均缺口1.08），"
        "验证了本章算法设计在该口径下的有效性。"
    )


def update_existing_text(doc: Document) -> None:
    for p in doc.paragraphs:
        text = p.text.strip()
        if "实验对比折线图如图5-25所示" in text:
            p.text = p.text.replace("实验对比折线图如图5-25所示", "补充实验结果如图5-25所示")
        if "实验对比图如图5-26所示" in text:
            p.text = p.text.replace("实验对比图如图5-26所示", "补充实验结果如图5-26所示")
        if text.startswith("图5-25"):
            p.text = "图5-25  RF与简单平均法绝对误差对比图"
        if text.startswith("图5-26"):
            p.text = "图5-26  动态安全库存与固定法对比图"


def main() -> None:
    if not INPUT_DOC.exists():
        raise FileNotFoundError(f"输入文件不存在: {INPUT_DOC}")

    generate_images()
    doc = Document(str(INPUT_DOC))

    update_existing_text(doc)
    ok_25 = replace_placeholder_with_image(doc, "图5-25", IMG_RF)
    ok_26 = replace_placeholder_with_image(doc, "图5-26", IMG_SS)
    if not ok_25 or not ok_26:
        raise RuntimeError(f"图占位替换失败: 图5-25={ok_25}, 图5-26={ok_26}")

    insert_ai9_section(doc)
    doc.save(str(OUTPUT_DOC))
    print(f"OK: {OUTPUT_DOC}")
    print(f"IMG: {IMG_RF}")
    print(f"IMG: {IMG_SS}")


if __name__ == "__main__":
    main()
